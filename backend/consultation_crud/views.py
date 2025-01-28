from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from .models import Consultation, Ordonnance, Medicament,ResumeConsultation,SoinsInfirmier

from  authentication.models import Dpis,Infirmiers,Medecins

from authentication.serializers import DpiSerializer, MedecinSerializer
from .serializers import ConsultationSerializer, OrdonnanceSerializer,MedicamentSerializer,ResumeConsultationSerializer,SoinsInfirmierSerializer,Consultation_Serializer,Ordonnance_Serializer
from django.shortcuts import get_object_or_404, redirect
from io import BytesIO
from django.http import HttpResponse
from django.db import transaction
from datetime import date
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.shared import Pt
import os
from django.conf import settings
import requests



@api_view(['GET'])
def get_ordonnance_id_by_consultation(request, consultation_id):
    try:
        # Récupérer l'ordonnance associée à la consultation
        ordonnance = Ordonnance.objects.get(consultation_id=consultation_id)
        
        # Retourner uniquement l'ID de l'ordonnance
        return Response({'ordonnance_id': ordonnance.id}, status=status.HTTP_200_OK)
    except Ordonnance.DoesNotExist:
        return Response({'error': 'Ordonnance not found for this consultation'}, status=status.HTTP_404_NOT_FOUND)

@api_view(['POST'])
def create_medicament(request):
    # Vérifiez que l'ordonnance existe et récupérez-la
    ordonnance_id = request.data.get('ordonnance_id')
    try:
        ordonnance = Ordonnance.objects.get(id=ordonnance_id)
    except Ordonnance.DoesNotExist:
        return Response({'error': 'Ordonnance not found'}, status=status.HTTP_400_BAD_REQUEST)

    # Créez le médicament avec l'ordonnance associée
    medicament = Medicament.objects.create(
        nom_medicament=request.data.get('nom_medicament'),
        dose=request.data.get('dose'),
        duree=request.data.get('duree'),
        ordonnance=ordonnance  # Assurez-vous que l'ordonnance est liée
    )

    # Réponse de succès
    return Response({'message': 'Medicament created successfully'}, status=status.HTTP_201_CREATED)

@api_view(['GET'])
def list_medicaments(request):
    medicaments = Medicament.objects.all()  # Récupérer tous les médicaments
    serializer = MedicamentSerializer(medicaments, many=True)  # Sérialiser la liste des médicaments
    return Response(serializer.data, status=status.HTTP_200_OK)

##details : 
@api_view(['GET'])
def detail_medicament(request, medicament_id):
    medicament = get_object_or_404(Medicament, id=medicament_id)  # Récupérer le médicament par son ID
    serializer = MedicamentSerializer(medicament)  # Sérialiser le médicament
    return Response(serializer.data, status=status.HTTP_200_OK)

##Update : 
@api_view(['PUT'])
def update_medicament(request, medicament_id):
    medicament = get_object_or_404(Medicament, id=medicament_id)  # Récupérer le médicament par son ID
    serializer = MedicamentSerializer(medicament, data=request.data, partial=True)  # Permet la mise à jour partielle
    
    if serializer.is_valid():
        serializer.save()  # Sauvegarder les modifications
        return Response(serializer.data, status=status.HTTP_200_OK)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


##Delete 
@api_view(['DELETE'])
def delete_medicament(request, medicament_id):
    medicament = get_object_or_404(Medicament, id=medicament_id)  # Récupérer le médicament à supprimer
    medicament.delete()  # Supprimer le médicament
    return Response(status=status.HTTP_204_NO_CONTENT)  # Retourner une réponse 204 (pas de contenu)






##   2.crud_consultation##
@api_view(['POST'])
def create_consultation(request):
    """
    Endpoint pour créer une consultation.
    """
    serializer = ConsultationSerializer(data=request.data)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


# Liste des consultations

@api_view(['GET'])
def list_consultations(request):
    patient_id = request.query_params.get('patient_id', None)
    consultations = Consultation.objects.filter(patient_id=patient_id) if patient_id else Consultation.objects.all()
    consultations = consultations.order_by('-date_consultation')  # Trier par date décroissante

    serializer = ConsultationSerializer(consultations, many=True)
    return Response(serializer.data, status=status.HTTP_200_OK)

@api_view(['GET'])
def list_consultations_by_medecin(request, medecins_id):
    """
    Liste les consultations associées à un médecin spécifique.
    """

    consultations = Consultation.objects.filter(patient__medecins__id=medecins_id).order_by('-date_consultation')
    if not consultations.exists():
        return Response({"message": "Aucune consultation trouvée pour ce médecin."}, status=status.HTTP_404_NOT_FOUND)
    
    serializer = ConsultationSerializer(consultations, many=True)
    return Response(serializer.data, status=status.HTTP_200_OK)

@api_view(['GET'])
def list_consultations_by_patient(request, nss):
    """
    Liste les consultations avec leurs ordonnances associées pour un patient spécifique.
    """
    consultations = Consultation.objects.filter(patient__nss=nss).order_by('-date_consultation')
    if not consultations.exists():
        return Response({"message": "Aucune consultation trouvée pour ce patient."}, status=status.HTTP_404_NOT_FOUND)
    
    # Sérialisation des consultations
    serializer = Consultation_Serializer(consultations, many=True)
    
    # Ajouter dynamiquement l'ID de l'ordonnance
    response_data = serializer.data
    for consultation in response_data:
        consultation_id = consultation['id']
        ordonnances = Ordonnance.objects.filter(consultation_id=consultation_id)
        
        # Si une ordonnance existe, on ajoute seulement son ID
        if ordonnances.exists():
            consultation['ordonnance'] = ordonnances.first().id  # Ajouter l'ID de l'ordonnance
        else:
            consultation['ordonnance'] = None  # Aucun ID d'ordonnance si elle n'existe pas
    

    

    
    return Response(response_data, status=status.HTTP_200_OK)


# Détails d'une consultation spécifique
@api_view(['GET'])
def detail_consultation(request, consultation_id):
    consultation = get_object_or_404(Consultation, id=consultation_id)
    serializer = ConsultationSerializer(consultation)
    return Response(serializer.data, status=status.HTTP_200_OK)

# Mise à jour d'une consultation spécifique
@api_view(['PUT'])
def update_consultation(request, consultation_id):
    consultation = get_object_or_404(Consultation, id=consultation_id)
    serializer = ConsultationSerializer(consultation, data=request.data, partial=True)  # Mise à jour partielle

    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=status.HTTP_200_OK)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

# Suppression d'une consultation
@api_view(['DELETE'])
def delete_consultation(request, consultation_id):
    consultation = get_object_or_404(Consultation, id=consultation_id)
    consultation.delete()
    return Response(status=status.HTTP_204_NO_CONTENT)



@api_view(['POST'])
def create_resume_consultation(request):
    if request.method == 'POST':
        # Sérialiser les données du résumé de consultation
        serializer = ResumeConsultationSerializer(data=request.data)

        if serializer.is_valid():
            # Récupérer l'ID de la consultation depuis la requête
            consultation_id = request.data.get('consultationId')

            try:
                # Trouver la consultation par son ID
                consultation = Consultation.objects.get(id=consultation_id)

                # Sauvegarder le résumé de consultation dans la base de données
                resume_consultation = serializer.save()

                # Associer le résumé de consultation à la consultation
                consultation.resume_consultation = resume_consultation
                consultation.save()

                # Répondre avec les données du résumé de consultation
                return Response(serializer.data, status=status.HTTP_201_CREATED)

            except Consultation.DoesNotExist:
                # Si la consultation n'existe pas, renvoyer une erreur
                return Response({"error": "Consultation not found"}, status=status.HTTP_400_BAD_REQUEST)

        # Si le serializer n'est pas valide, retourner une erreur
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


# Liste des résumés de consultations
@api_view(['GET'])
def list_resume_consultations(request):
    resume_consultations = ResumeConsultation.objects.all()
    serializer = ResumeConsultationSerializer(resume_consultations, many=True)
    return Response(serializer.data, status=status.HTTP_200_OK)



@api_view(['GET'])
def detail_resume_consultation(request, consultation_id):
    # Récupérer la consultation à partir de l'ID de la consultation
    consultation = get_object_or_404(Consultation, id=consultation_id)

    # Récupérer le résumé de consultation associé
    resume_consultation = consultation.resume_consultation
    if resume_consultation:
        serializer = ResumeConsultationSerializer(resume_consultation)
        return Response(serializer.data, status=status.HTTP_200_OK)
    else:
        return Response({'detail': 'Aucun résumé de consultation trouvé pour cette consultation.'}, 
                        status=status.HTTP_404_NOT_FOUND)

# Mise à jour d'un résumé de consultation
@api_view(['PUT'])
def update_resume_consultation(request, resume_consultation_id):
    resume_consultation = get_object_or_404(ResumeConsultation, id=resume_consultation_id)
    serializer = ResumeConsultationSerializer(resume_consultation, data=request.data, partial=True)

    if serializer.is_valid():
        # Sauvegarder le résumé mis à jour
        resume_consultation = serializer.save()

        # Associer ce résumé à la consultation
        consultation_id = request.data.get('consultation_id', None)
        if consultation_id:
            consultation = get_object_or_404(Consultation, id=consultation_id)
            consultation.resume_consultation = resume_consultation
            consultation.save()

        return Response(serializer.data, status=status.HTTP_200_OK)

    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


# Suppression d'un résumé de consultation
@api_view(['DELETE'])
def delete_resume_consultation(request, resume_consultation_id):
    resume_consultation = get_object_or_404(ResumeConsultation, id=resume_consultation_id)
    resume_consultation.delete()
    return Response(status=status.HTTP_204_NO_CONTENT)

@api_view(['POST'])
def create_ordonnance(request, consultation_id):
    # Récupérer la consultation
    consultation = get_object_or_404(Consultation, id=consultation_id)

    # Vérifier si une ordonnance existe déjà pour cette consultation
    if Ordonnance.objects.filter(consultation=consultation).exists():
        return Response({"message": "Une ordonnance existe déjà pour cette consultation."}, status=400)

    # Créer l'ordonnance
    ordonnance_data = {
        'consultation': consultation.id,
        'etat_ordonnance': 'En cours de validation'
    }
    ordonnance_serializer = OrdonnanceSerializer(data=ordonnance_data)
    
    if ordonnance_serializer.is_valid():
        ordonnance = ordonnance_serializer.save()

        # Retourner l'ordonnance créée
        return Response(ordonnance_serializer.data, status=status.HTTP_201_CREATED)
    
    return Response(ordonnance_serializer.errors, status=status.HTTP_400_BAD_REQUEST)

@api_view(['GET'])
def get_ordonnance(request, ordonnance_id):
    # Récupérer l'ordonnance à partir de l'ID
    ordonnance = get_object_or_404(Ordonnance, id=ordonnance_id)

    # Sérialiser les données de l'ordonnance
    serializer = OrdonnanceSerializer(ordonnance)

    # Retourner la réponse avec les données de l'ordonnance
    return Response(serializer.data)


@api_view(['PUT'])
def update_ordonnance(request, ordonnance_id):
    # Récupérer l'ordonnance via l'ID
    ordonnance = get_object_or_404(Ordonnance, id=ordonnance_id)

    # Mettre à jour l'ordonnance avec les nouvelles données
    ordonnance_serializer = OrdonnanceSerializer(ordonnance, data=request.data, partial=True)

    if ordonnance_serializer.is_valid():
        ordonnance = ordonnance_serializer.save()

        # Si des médicaments sont fournis, les mettre à jour
        medicaments_data = request.data.get('medicaments', [])
        if medicaments_data:
            # Supprimer les anciens médicaments associés à l'ordonnance avant d'ajouter les nouveaux
            ordonnance.medicaments.all().delete()

            for medicament_data in medicaments_data:
                medicament_data['ordonnance'] = ordonnance.id  # Associer l'ordonnance au médicament
                medicament_serializer = MedicamentSerializer(data=medicament_data)
                if medicament_serializer.is_valid():
                    # Sauvegarder chaque médicament
                    medicament_serializer.save()
                else:
                    # Enregistrer les erreurs de validation des médicaments
                    return Response(medicament_serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        # Retourner les données de l'ordonnance mise à jour avec les médicaments
        return Response(ordonnance_serializer.data, status=status.HTTP_200_OK)

    # Si l'ordonnance n'est pas valide, retourner les erreurs
    return Response(ordonnance_serializer.errors, status=status.HTTP_400_BAD_REQUEST)
@api_view(['DELETE'])
def delete_ordonnance(request, ordonnance_id):
    ordonnance = get_object_or_404(Ordonnance, id=ordonnance_id)
    ordonnance.delete()
    return Response(status=status.HTTP_204_NO_CONTENT)

@api_view(['GET'])
def download_ordonnance_word(request, ordonnance_id):
    # Récupérer l'ordonnance, la consultation et les informations associées
    ordonnance = get_object_or_404(Ordonnance, id=ordonnance_id)
    consultation = ordonnance.consultation
    patient = consultation.patient

    # Accéder au médecin associé au patient
    medecin = patient.medecins
    if medecin:
        # Récupérer l'objet 'Personnel' du médecin
        personnel_medecin = medecin.personnel
        medecin_nom = personnel_medecin.nom
        medecin_prenom = personnel_medecin.prenom
        medecin_specialite = medecin.specialite or 'Médecin généraliste'
    else:
        # Si le médecin est inconnu, définir des valeurs par défaut
        medecin_nom = 'Inconnu'
        medecin_prenom = ''
        medecin_specialite = 'Médecin inconnu'

    # Création du document Word
    doc = Document()

    # Ajouter un tableau pour les informations du patient et du médecin
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    row = table.rows[0].cells

    # Informations du patient dans la colonne gauche
    patient_info = row[0].paragraphs[0]
    patient_info.add_run(f"Nom : {patient.nom}\n").bold = True
    patient_info.add_run(f"Prénom : {patient.prenom}\n")
    patient_info.add_run(f"Adresse : {patient.adresse or 'Non renseignée'}\n")

    # Informations du médecin dans la colonne droite
    doctor_info = row[1].paragraphs[0]
    doctor_info.add_run(f"Dr {medecin_nom} {medecin_prenom}\n").bold = True
    doctor_info.add_run(f"{medecin_specialite}\n")

    # Alignements et espacements
    for cell in row:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doctor_info.paragraph_format.left_indent = Pt(50)

    # Ajouter la date de consultation en gras noir
    date_paragraph = doc.add_paragraph()
    date_paragraph.add_run(f"Consulté le: {consultation.date_consultation.strftime('%d/%m/%Y') if consultation.date_consultation else 'Date non renseignée'}").bold = False
    doc.add_paragraph("\n")
    
    # Ajouter le titre "Ordonnance"
    title = doc.add_paragraph()
    title_run = title.add_run("Ordonnance")
    title_run.bold = True
    title_run.font.size = Pt(20)  # Grande taille
    title_run.font.color.rgb = None  # Couleur noire
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Éviter toute bordure ou espace supplémentaire
    title.paragraph_format.space_after = Pt(0)  # Pas d'espace après
    title.paragraph_format.space_before = Pt(0)  # Pas d'espace avant

    # Ajouter un saut de ligne avant les médicaments
    doc.add_paragraph("\n\n")  # Sauts de ligne avant la liste des médicaments

    # Ajouter les médicaments
    for idx, medicament in enumerate(ordonnance.medicaments.all(), start=1):
        p = doc.add_paragraph()
        p.add_run(f"{idx}. {medicament.nom_medicament} ").bold = False
        p.add_run(f"{medicament.dose} ").bold = True
        p.add_run("pendant ").bold = False
        p.add_run(f"{medicament.duree} jours").bold = True
        doc.add_paragraph()  # Ajouter une ligne vide entre les médicaments

    # Ajouter la signature en gras, alignée à droite
    signature_paragraph = doc.add_paragraph()
    signature_paragraph.add_run("\nSignature :").bold = True
    signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Sauvegarder le document et le retourner en réponse HTTP
    filename = f"{patient.nom}_{patient.prenom}_{consultation.date_consultation.strftime('%d%m%Y') if consultation.date_consultation else 'unknown_date'}.docx"
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    doc.save(response)

    return response

@api_view(['GET'])
def download_resume_consultation_word(request, consultation_id):
    # Récupérer la consultation à partir de l'ID
    consultation = get_object_or_404(Consultation, id=consultation_id)
    
    # Récupérer le résumé de consultation
    resume = consultation.resume_consultation
    if not resume:
        return HttpResponse("Aucun résumé de consultation disponible", status=404)
    
    # Création du document Word
    doc = Document()

    # Ajouter les informations de consultation dans le document
    doc.add_heading(f"Résumé de Consultation", 0)

    # Informations du patient
    doc.add_heading("Informations du patient", level=1)
    doc.add_paragraph(f"Nom : {consultation.patient.nom}")
    doc.add_paragraph(f"Prénom : {consultation.patient.prenom}")
    doc.add_paragraph(f"NSS : {consultation.patient.nss}")
    
    # Informations de consultation
    doc.add_heading("Détails de la consultation", level=1)
    doc.add_paragraph(f"Date de la consultation : {consultation.date_consultation.strftime('%d/%m/%Y')}")
    
    # Résumé de la consultation
    doc.add_heading("Résumé de consultation", level=1)
    doc.add_paragraph(f"Diagnostic : {resume.diagnostic}")
    doc.add_paragraph(f"Symptômes : {resume.symptomes}")
    doc.add_paragraph(f"Mesures prises : {resume.mesure}")
    doc.add_paragraph(f"Date de prochaine consultation : {resume.date_prochaine_consultation.strftime('%d/%m/%Y') if resume.date_prochaine_consultation else 'Non renseignée'}")
    
    # Ajouter des informations supplémentaires si disponibles
    if resume.antecedents:
        doc.add_paragraph(f"Antécédents : {resume.antecedents}")
    if resume.info_supp:
        doc.add_paragraph(f"Informations supplémentaires : {resume.info_supp}")

    # Sauvegarder le document et le retourner en réponse HTTP
    filename = f"{consultation.patient.nom}_{consultation.patient.prenom}_resume_{consultation.date_consultation.strftime('%d%m%Y')}.docx"
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    doc.save(response)

    return response
@api_view(['POST'])
def valider_ordonnance(request, ordonnance_id):
    try:
        ordonnance = Ordonnance.objects.get(id=ordonnance_id)
    except Ordonnance.DoesNotExist:
        return Response({"error": "Ordonnance non trouvée"}, status=status.HTTP_404_NOT_FOUND)

    # Paramètres à envoyer au SGPH
    data = {
        'ordonnance_id': ordonnance.id,
        'medicaments': [
            {'nom_medicament': medicament.nom_medicament, 'dose': medicament.dose, 'duree': medicament.duree}
            for medicament in ordonnance.medicaments.all()
        ]
    }

    # URL de l'API SGPH pour valider l'ordonnance
    sgph_url = "http://localhost:8001/api/valider_ordonnance/"  # Adaptez l'URL selon votre configuration

    try:
        response = requests.post(sgph_url, json=data)  # Envoi de la requête POST au SGPH

        if response.status_code == 200:
            # Traitement de la réponse du SGPH
            sgph_response = response.json()
            etat_ordonnance = sgph_response.get('etat_ordonnance', 'En cours de validation')  # Valeur par défaut
            motif_refus = sgph_response.get('motif_refus', None)

            # Vérification des valeurs valides pour 'etat_ordonnance'
            if etat_ordonnance not in ['Validee', 'En cours de validation', 'Refusee']:
                return Response({
                    "error": "État de l'ordonnance invalide reçu du SGPH"
                }, status=status.HTTP_400_BAD_REQUEST)

            # Mise à jour de l'état de l'ordonnance dans le DPI
            ordonnance.etat_ordonnance = etat_ordonnance
            ordonnance.motif_refus = motif_refus
            ordonnance.save()

            # Retour de la réponse avec un message confirmant la mise à jour
            return Response({
                "message": f"Ordonnance {etat_ordonnance} avec succès.",
                "motif_refus": motif_refus if motif_refus else "Aucun motif de refus"
            }, status=status.HTTP_200_OK)

        else:
            return Response({
                "error": "Échec de la validation dans le SGPH",
                "details": response.json()  # Inclut les détails du message d'erreur si présents
            }, status=status.HTTP_400_BAD_REQUEST)

    except requests.exceptions.RequestException as e:
        # Gestion des erreurs liées à la requête HTTP
        return Response({
            "error": "Erreur de communication avec le SGPH",
            "details": str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    




@api_view(['GET'])
def list_ordonnances_by_medecin(request, medecins_id):
    """
    Liste les ordonnances associées à un médecin spécifique.
    """
    # Récupérer les consultations
    consultations = Consultation.objects.filter(patient__medecins__id=medecins_id)

    if not consultations.exists():
        return Response({"message": "Aucune consultation trouvée pour ce médecin."}, status=status.HTTP_404_NOT_FOUND)

    # Récupérer les ordonnances
    ordonnances = Ordonnance.objects.filter(consultation__in=consultations)

    if not ordonnances.exists():
        return Response({"message": "Aucune ordonnance trouvée pour ce médecin."}, status=status.HTTP_404_NOT_FOUND)

    # Sérialiser les ordonnances
    serializer = OrdonnanceSerializer(ordonnances, many=True)

    # Retourner les données sérialisées
    return Response(serializer.data, status=status.HTTP_200_OK)
@api_view(['GET'])
def list_ordonnances_by_patient(request, patient_id):
    """
    Liste les ordonnances associées à un patient spécifique.
    """
    consultations = Consultation.objects.filter(patient_id=patient_id)
    if not consultations.exists():
        return Response({"message": "Aucune consultation trouvée pour ce patient."}, status=status.HTTP_404_NOT_FOUND)
    
    ordonnances = Ordonnance.objects.filter(consultation__in=consultations)
    if not ordonnances.exists():
        return Response({"message": "Aucune ordonnance trouvée pour ce patient."}, status=status.HTTP_404_NOT_FOUND)
    
    serializer = OrdonnanceSerializer(ordonnances, many=True)
    return Response(serializer.data, status=status.HTTP_200_OK)


## crud soins_infermier ##

@api_view(['POST'])
def create_soin(request):
    if request.method == 'POST':
        serializer = SoinsInfirmierSerializer(data=request.data, context={'request': request})  # Passer `request`
        if serializer.is_valid():
            soin = serializer.save()  # Sauvegarde en base
            
            # Ajouter `heure_soin` et `type_soin` à la réponse
            response_data = serializer.data
            response_data['heure_soin'] = request.data.get('heure_soin', "00:00")
            response_data['type_soin'] = request.data.get('type_soin', "Type inconnu")

            return Response(response_data, status=status.HTTP_201_CREATED)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

# Liste des soins
@api_view(['GET'])
def list_soins(request):
    soins = SoinsInfirmier.objects.all()
    serializer = SoinsInfirmierSerializer(soins, many=True)
    return Response(serializer.data, status=status.HTTP_200_OK)

# Détails d'un soin spécifique
@api_view(['GET'])
def detail_soin(request, SoinsInfirmier_id):
    soin = get_object_or_404(SoinsInfirmier, id=SoinsInfirmier_id)
    serializer = SoinsInfirmierSerializer(soin)  # Passer l'instance de soin ici
    return Response(serializer.data, status=status.HTTP_200_OK)

# Mise à jour d'un soin spécifique
@api_view(['PUT'])
def update_soin(request, SoinsInfirmier_id):
    soin = get_object_or_404(SoinsInfirmier, id=SoinsInfirmier_id)
    serializer = SoinsInfirmierSerializer(soin, data=request.data, partial=True)  # Mise à jour partielle
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=status.HTTP_200_OK)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

# Suppression d'un soin
@api_view(['DELETE'])
def delete_soin(request, SoinsInfirmier_id):
    soin = get_object_or_404(SoinsInfirmier, id=SoinsInfirmier_id)
    soin.delete()
    return Response(status=status.HTTP_204_NO_CONTENT)


@api_view(['GET'])
def get_patient_id_by_nss(request, nss):
    """
    Récupère l'ID du patient à partir du NSS.
    """
    try:
        # Recherche du patient par NSS
        patient = Dpis.objects.get(nss=nss)
        return Response({"patient_id": patient.id}, status=status.HTTP_200_OK)
    except Dpis.DoesNotExist:
        return Response({"message": "Aucun patient trouvé avec ce NSS."}, status=status.HTTP_404_NOT_FOUND)